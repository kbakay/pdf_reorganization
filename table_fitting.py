from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import re
import os

def set_cell_text(cell, text):
    cell.text = text
    cell._element.get_or_add_tcPr().append(OxmlElement('w:vAlign'))
    cell._element.tcPr.get_or_add_vAlign().set(qn('w:val'), "center")

def set_row_height(row, height):
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    trHeight = OxmlElement('w:trHeight')
    trHeight.set(qn('w:val'), str(height))
    trHeight.set(qn('w:hRule'), 'atLeast')
    trPr.append(trHeight)

def remove_page_numbers(document):
    pattern = re.compile(r'^\d+/\d+$')
    for paragraph in document.paragraphs:
        if pattern.match(paragraph.text.strip()):
            paragraph.text = ""

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if pattern.match(cell.text.strip()):
                    cell.text = ""

def connect_table_cells(document):
    tables = document.tables
    for table_index in range(1, len(tables)):
        current_table = tables[table_index]
        previous_table = tables[table_index - 1]

        current_table_first_cell_text = current_table.rows[0].cells[0].text.strip()
        if not current_table_first_cell_text:
            previous_table_last_row_index = len(previous_table.rows) - 1
            if previous_table_last_row_index >= 0:
                previous_table_last_row = previous_table.rows[previous_table_last_row_index]
                for row_index, row in enumerate(current_table.rows):
                    for cell_index, cell in enumerate(row.cells):
                        previous_text = previous_table_last_row.cells[cell_index].text.strip()
                        current_text = cell.text.strip()
                        if (not current_text or previous_text) and previous_table_last_row.cells[cell_index].text.strip():
                            combined_text = f"{previous_text} {current_text}".strip()
                            set_cell_text(cell, combined_text)
                            previous_table_last_row.cells[cell_index].text = ""

    for table in tables:
        rows_to_delete = []
        for row_index, row in enumerate(table.rows):
            if all(cell.text.strip() == "" for cell in row.cells):
                rows_to_delete.append(row_index)

        for row_index in reversed(rows_to_delete):
            table._tbl.remove(table.rows[row_index]._tr)

    # Ayar satır yüksekliği
    for table in tables:
        for row in table.rows:
            max_len = max(len(cell.text) for cell in row.cells)
            if max_len < 50:
                height = 300
            elif max_len < 100:
                height = 600
            elif max_len < 200:
                height = 900
            else:
                height = 1200
            set_row_height(row, height)

    # Sayfa numaralarını kaldır
    remove_page_numbers(document)

input_directory = "3"
output_directory = "4"

# Kontrol etme ve oluşturma
if not os.path.exists(output_directory):
    os.makedirs(output_directory)

# Input dosyalarını al ve işlem yap
for filename in os.listdir(input_directory):
    if filename.endswith(".docx"):
        input_path = os.path.join(input_directory, filename)
        output_path = os.path.join(output_directory, filename)
        document = Document(input_path)
        connect_table_cells(document)
        document.save(output_path)


