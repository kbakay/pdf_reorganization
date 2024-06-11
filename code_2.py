from docx import Document
import os

def process_cell_text(cell):
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            if run.font.strike:
                # Mevcut paragrafların başına **/** ekle
                new_run = paragraph.insert_paragraph_before().add_run("**/** ")
                new_run.font.strike = run.font.strike
                return

def process_table(table):
    for row in table.rows:
        process_cell_text(row.cells[0])

def process_documents(input_folder, output_folder):
    # Çıktı klasörünü kontrol etme ve oluşturma
    if not os.path.exists(output_folder):
        os.makedirs(output_folder)

    # Dosyaları işleme ve kaydetme
    for filename in os.listdir(input_folder):
        if filename.endswith("updated.docx"):
            input_path = os.path.join(input_folder, filename)
            output_path = os.path.join(output_folder, filename)
            
            # Word dosyasını yükleyin
            doc = Document(input_path)

            # Tabloları döngüye alarak her tabloyu işleyin
            for table in doc.tables:
                # İlk hücrenin içeriğini kontrol ederek "ipotek" kısmına ulaşıp ulaşmadığınızı kontrol edin
                first_cell_text = table.rows[0].cells[0].text.strip().lower()
                if "ipotek" in first_cell_text:
                    # İpotek tablosu için özel işlem (hücre içeriğini değiştirmeden sadece başına ekleyin)
                    for row in table.rows:
                        process_cell_text(row.cells[0])
                else:
                    process_table(table)

            # Değişiklikleri kaydedin
            doc.save(output_path)

if __name__ == "__main__":
    input_folder = "1"
    output_folder = "2"

    process_documents(input_folder, output_folder)
