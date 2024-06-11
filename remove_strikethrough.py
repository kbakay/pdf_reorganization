from docx import Document
import os

def remove_strikethrough_from_paragraph(paragraph):
    """Bir paragraftaki tüm üstü çizili yazıların çizgisini kaldırır."""
    for run in paragraph.runs:
        if run.font.strike:
            run.font.strike = False

def remove_strikethrough_from_cell(cell):
    """Bir hücredeki tüm paragraflardaki üstü çizili yazıların çizgisini kaldırır."""
    for paragraph in cell.paragraphs:
        remove_strikethrough_from_paragraph(paragraph)
    # Herhangi bir nested tablo varsa onu da işlemeliyiz.
    for table in cell.tables:
        remove_strikethrough_from_table(table)

def remove_strikethrough_from_table(table):
    """Bir tablodaki tüm hücrelerdeki üstü çizili yazıların çizgisini kaldırır."""
    for row in table.rows:
        for cell in row.cells:
            remove_strikethrough_from_cell(cell)

def remove_strikethrough(file_path, output_path):
    doc = Document(file_path)
    
    # Belge içindeki tüm tabloları işlemek için
    for table in doc.tables:
        remove_strikethrough_from_table(table)
    
    # Belge içindeki ana paragrafları işlemek için
    for paragraph in doc.paragraphs:
        remove_strikethrough_from_paragraph(paragraph)

    doc.save(output_path)

def batch_remove_strikethrough(input_folder, output_folder):
    # Çıktı klasörünü kontrol etme ve oluşturma
    if not os.path.exists(output_folder):
        os.makedirs(output_folder)

    # Dosyaları işleme ve kaydetme
    for filename in os.listdir(input_folder):
        if filename.endswith(".docx"):
            input_path = os.path.join(input_folder, filename)
            output_path = os.path.join(output_folder, filename)
            
            remove_strikethrough(input_path, output_path)

    print("İşlem tamamlandı.")

if __name__ == "__main__":
    input_folder = "2"
    output_folder = "3"

    batch_remove_strikethrough(input_folder, output_folder)
