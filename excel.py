from docx import Document
import pandas as pd
import os

def extract_tables_from_word(file_path):
    # Open the Word document
    doc = Document(file_path)
    tables = []
    table_count = 0
    # Iterate through each table in the Word document
    for table in doc.tables:
        table_count += 1
        if table_count <= 2:  # Skip first two tables
            continue
        data = []
        # Iterate through each row in the table
        for row in table.rows:
            row_data = []
            # Iterate through each cell in the row
            for cell in row.cells:
                # Recursive function to extract tables from each cell
                cell_tables = extract_tables_from_cell(cell)
                if cell_tables:
                    for sub_table in cell_tables:
                        row_data.extend(sub_table)
                else:
                    row_data.append(cell.text)
            data.append(row_data)
        tables.append(pd.DataFrame(data))
    return tables

def extract_tables_from_cell(cell):
    tables = []
    # Check for tables in the cell
    for table in cell.tables:
        table_data = []
        # Iterate through each row in the table
        for row in table.rows:
            row_data = []
            # Iterate through each cell in the row
            for cell in row.cells:
                # Recursive call to extract tables from cell
                cell_tables = extract_tables_from_cell(cell)
                if cell_tables:
                    for sub_table in cell_tables:
                        row_data.extend(sub_table)
                else:
                    row_data.append(cell.text)
            table_data.append(row_data)
        tables.append(table_data)
    return tables

def save_tables_to_excel(tables, excel_file):
    # Concatenate all tables into a single DataFrame
    combined_table = pd.concat(tables, ignore_index=True)
    # Write to Excel file
    combined_table.to_excel(excel_file, index=False)
    print("Exported.")

def batch_process_word_to_excel(input_folder, output_folder):
    # Check and create if the output folder doesn't exist
    if not os.path.exists(output_folder):
        os.makedirs(output_folder)

    # Process and save each file in the input folder
    for filename in os.listdir(input_folder):
        if filename.endswith("updated.docx"):
            input_path = os.path.join(input_folder, filename)
            output_path = os.path.join(output_folder, os.path.splitext(filename)[0] + ".xlsx")
            # Extract tables from the Word file
            tables = extract_tables_from_word(input_path)
            # Save tables to Excel
            save_tables_to_excel(tables, output_path)

    print("Process completed.")

if __name__ == "__main__":
    input_folder = "4"
    output_folder = "output_excel"
    batch_process_word_to_excel(input_folder, output_folder)
